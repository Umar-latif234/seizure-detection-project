"""
generate_paper.py — IEEE-format research paper generator.

Produces: research_paper/Seizure_Detection_Research_Paper.docx
"""

import os
import io
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
ROOT       = os.path.dirname(os.path.abspath(__file__))
RESULTS    = os.path.join(ROOT, "results")
OUT_DIR    = os.path.join(ROOT, "research_paper")
OUT_FILE   = os.path.join(OUT_DIR, "Seizure_Detection_Research_Paper.docx")
CSV_PATH   = os.path.join(RESULTS, "comparison_results.csv")

os.makedirs(OUT_DIR, exist_ok=True)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=10, bold=False,
             italic=False, color=None):
    run.font.name      = name
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=10, bold=False, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment        = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level_label, size=11, space_before=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"{level_label}  {text.upper()}")
    set_font(run, size=size, bold=True)
    return p


def add_subheading(doc, text, space_before=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, bold=True, italic=True)
    return p


def add_figure(doc, img_path, caption, width=6.0):
    if not os.path.exists(img_path):
        add_paragraph(doc, f"[Figure not found: {img_path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, size=9, italic=True)


def add_ieee_table(doc, headers, rows, caption, col_widths=None):
    """Add an IEEE-style table (horizontal rules only, no vertical lines)."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # Remove all borders, then add only top/bottom/header-bottom
    def set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = OxmlElement(f"w:{edge}")
            if edge in kwargs:
                tag.set(qn("w:val"),   kwargs[edge].get("val",   "single"))
                tag.set(qn("w:sz"),    kwargs[edge].get("sz",    "4"))
                tag.set(qn("w:space"), kwargs[edge].get("space", "0"))
                tag.set(qn("w:color"), kwargs[edge].get("color", "000000"))
            else:
                tag.set(qn("w:val"), "none")
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(h)
        set_font(run, size=9, bold=True)
        set_cell_border(cell,
                        top={"val": "single", "sz": "6"},
                        bottom={"val": "single", "sz": "4"})

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        is_last = (r_idx == len(rows) - 1)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run  = p.add_run(str(val))
            set_font(run, size=9)
            border_args = {}
            if is_last:
                border_args["bottom"] = {"val": "single", "sz": "6"}
            set_cell_border(cell, **border_args)

    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(8)
    r = cap.add_run(caption)
    set_font(r, size=9, italic=True)


def add_ruled_line(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


# ─────────────────────────────────────────────────────────────────────────────
# LOAD RESULTS
# ─────────────────────────────────────────────────────────────────────────────
df = pd.read_csv(CSV_PATH)

uci     = df[df["dataset"] == "UCI"]
chb     = df[df["dataset"] == "CHB-MIT"]

# Best configurations
uci_best  = uci.sort_values("f1_score", ascending=False).iloc[0]
chb_best  = chb.sort_values("pr_auc",   ascending=False).iloc[0]

# Pipeline comparison means
pipe_means = df.groupby(["dataset", "pipeline"])[["accuracy", "f1_score", "pr_auc"]].mean()

# Penalty means across datasets
pen_means  = df.groupby("penalty")[["f1_score", "pr_auc"]].mean()

# Sparsity from logs (hardcoded from run output)
SPARSITY = {"L1 (Lasso)": (46, 178, 74.2), "ElasticNet": (109, 178, 38.8), "L2 (Ridge)": (178, 178, 0.0)}

# Scenario table from analysis
SCENARIOS = [
    ("Underfitting  (C=0.001, k=5)",  0.001,  5,  0.369, 0.369),
    ("Normal        (C=1.0,   k=50)", 1.0,   50,  0.375, 0.370),
    ("Overfitting   (C=1000,  k=178)",1000, 178,  0.990, 0.366),
]

# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins (IEEE: 0.75" sides, 1" top/bottom)
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)


# ── TITLE ─────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(6)
r = title_p.add_run(
    "Epileptic Seizure Detection Using Logistic Regression:\n"
    "A Comparative Study of Preprocessing Pipelines, Regularisation,\n"
    "and Class Imbalance Handling"
)
set_font(r, size=16, bold=True)

# ── AUTHORS ───────────────────────────────────────────────────────────────────
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(4)
r = auth_p.add_run("Hashir Khan")
set_font(r, size=11, bold=True)

affil_p = doc.add_paragraph()
affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil_p.paragraph_format.space_after = Pt(2)
r = affil_p.add_run("Department of Computer Science")
set_font(r, size=10, italic=True)

email_p = doc.add_paragraph()
email_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
email_p.paragraph_format.space_after = Pt(10)
r = email_p.add_run("huzikhannn@gmail.com")
set_font(r, size=10, italic=True)

add_ruled_line(doc)

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_p.paragraph_format.space_before = Pt(6)
abs_p.paragraph_format.space_after  = Pt(4)
r = abs_p.add_run("Abstract—")
set_font(r, size=10, bold=True, italic=True)
r2 = abs_p.add_run(
    "Epileptic seizure detection from electroencephalography (EEG) signals is a critical "
    "clinical task hindered by extreme class imbalance and high-dimensional time-series data. "
    "This paper presents a systematic end-to-end machine learning study using Logistic Regression "
    "as the baseline classifier across two publicly available EEG datasets: the UCI Epileptic "
    "Seizure Recognition dataset and a subset of the CHB-MIT Scalp EEG dataset. "
    "We evaluate two distinct preprocessing pipelines—Pipeline A (normalisation, bandpass "
    "filtering, and univariate feature selection) and Pipeline B (statistical feature extraction, "
    "standard scaling, and PCA)—across three regularisation strategies (L1/Lasso, L2/Ridge, "
    "and ElasticNet) and three class-imbalance handling techniques (SMOTE oversampling, "
    "random undersampling, and class weighting), yielding 36 experiment configurations. "
    "Our results demonstrate that preprocessing order is the dominant performance factor: "
    "Pipeline B achieves a mean F1-score of 0.895 on UCI versus 0.371 for Pipeline A on the "
    "same data. Sparsity analysis confirms that L1 regularisation selects only 46 of 178 "
    "features (74.2\\% sparsity) compared to L2's dense solution. On the severely imbalanced "
    "CHB-MIT dataset (270:1 ratio), random undersampling achieves the highest PR-AUC of 0.625, "
    "outperforming SMOTE (0.528) and class weighting (0.471). All code, models, and results "
    "are fully reproducible."
)
set_font(r2, size=10)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.paragraph_format.space_after = Pt(8)
r = kw_p.add_run("Keywords—")
set_font(r, size=10, bold=True, italic=True)
r2 = kw_p.add_run(
    "epileptic seizure detection; EEG; logistic regression; preprocessing pipeline; "
    "regularisation; class imbalance; SMOTE; undersampling; PCA; feature selection."
)
set_font(r2, size=10, italic=True)

add_ruled_line(doc)

# ─────────────────────────────────────────────────────────────────────────────
# I. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Introduction", "I.")
add_paragraph(doc,
    "Epilepsy affects approximately 50 million people worldwide, making it one of the most "
    "prevalent neurological disorders [1]. Accurate and timely detection of epileptic seizures "
    "is essential for patient safety and treatment planning. Electroencephalography (EEG) remains "
    "the gold standard for seizure monitoring, but manual interpretation is time-consuming, "
    "expensive, and prone to inter-rater variability [2].")
add_paragraph(doc,
    "Automated seizure detection using machine learning offers a scalable alternative. However, "
    "several challenges make this a non-trivial problem: (1) extreme class imbalance—seizure "
    "segments represent less than 1\\% of recordings in real-world datasets; (2) high "
    "dimensionality of raw EEG time-series; and (3) strong sensitivity to preprocessing choices.")
add_paragraph(doc,
    "This work makes the following contributions: (i) a rigorous comparison of two fundamentally "
    "different preprocessing pipelines on two EEG datasets; (ii) a structured regularisation "
    "study including sparsity analysis; (iii) a comparison of three class-imbalance handling "
    "techniques; and (iv) explicit demonstration of underfitting and overfitting regimes using "
    "regularisation paths and learning curves. We deliberately restrict our model to Logistic "
    "Regression—a linear, interpretable classifier—to isolate the effects of preprocessing "
    "and regularisation from model complexity.")

# ─────────────────────────────────────────────────────────────────────────────
# II. DATASETS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Dataset Description", "II.")
add_paragraph(doc,
    "Three datasets were selected to cover different EEG recording conditions, imbalance "
    "levels, and feature characteristics. The Kaggle iEEG (Mayo Clinic) dataset was excluded "
    "from experiments due to API authentication requirements but is documented below.")

add_subheading(doc, "A.  UCI Epileptic Seizure Recognition Dataset")
add_paragraph(doc,
    "The UCI dataset [3] contains 11,500 samples derived from EEG recordings of 500 subjects. "
    "Each sample is a one-second EEG segment represented as 178 time-domain amplitude values "
    "sampled at approximately 173.6 Hz. The original five-class labels (seizure, tumour, "
    "healthy eyes open/closed, epileptic area without seizure) are binarised: class 1 "
    "(seizure) maps to positive; classes 2–5 map to negative.")

# Dataset characteristics table
add_ieee_table(doc,
    ["Property", "Value"],
    [
        ["Total samples",       "11,500"],
        ["Feature dimensionality", "178 (1-second EEG time series)"],
        ["Seizure samples (minority)", "2,300  (20.0%)"],
        ["Non-seizure samples (majority)", "9,200  (80.0%)"],
        ["Class imbalance ratio", "4 : 1"],
        ["Feature type",        "Raw time-series amplitudes"],
        ["Sampling frequency",  "~173.6 Hz"],
    ],
    "TABLE I.  UCI Dataset Characteristics",
    col_widths=[2.8, 3.5],
)

add_subheading(doc, "B.  CHB-MIT Scalp EEG Dataset (Subset)")
add_paragraph(doc,
    "The CHB-MIT dataset [4] contains long-term scalp EEG recordings from paediatric patients "
    "with intractable seizures. To respect a 40 GB storage constraint, only three one-hour "
    "recordings from Patient 1 (chb01_01.edf, chb01_02.edf, chb01_03.edf) are used. "
    "Files are segmented into non-overlapping 2-second windows (23 channels × 256 Hz × 2 s "
    "= 11,776 features per window). Seizure labels are derived from the accompanying summary "
    "file; windows overlapping annotated seizure intervals are labelled positive.")

add_ieee_table(doc,
    ["Property", "Value"],
    [
        ["Total windows",       "5,400"],
        ["Feature dimensionality", "11,776  (23 channels × 512 samples)"],
        ["Seizure windows (minority)", "20  (0.37%)"],
        ["Non-seizure windows (majority)", "5,380  (99.63%)"],
        ["Class imbalance ratio", "269 : 1"],
        ["Feature type",        "Raw multichannel EEG"],
        ["Sampling frequency",  "256 Hz"],
    ],
    "TABLE II.  CHB-MIT Dataset Characteristics",
    col_widths=[2.8, 3.5],
)

add_subheading(doc, "C.  Kaggle iEEG (Mayo Clinic) Dataset")
add_paragraph(doc,
    "The Kaggle Seizure Detection competition dataset [5] contains intracranial EEG (iEEG) "
    "recordings from dogs and humans, provided as MATLAB .mat files. Preictal segments "
    "(within one hour of seizure onset) are labelled positive; interictal segments (at least "
    "one week from any seizure) are labelled negative. Due to Kaggle API authentication "
    "requirements, this dataset could not be downloaded automatically and is excluded from "
    "experimental evaluation. Its inclusion in future work is recommended for cross-modal "
    "validation (scalp EEG vs. iEEG).")

# ─────────────────────────────────────────────────────────────────────────────
# III. PREPROCESSING PIPELINES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Preprocessing Pipeline Design", "III.")
add_paragraph(doc,
    "A central research question in this work is whether the ordering of preprocessing steps "
    "affects downstream classification performance. We design two structurally distinct "
    "pipelines implemented as scikit-learn Pipeline objects to ensure no data leakage "
    "between training and test sets.")

add_subheading(doc, "A.  Pipeline A: Signal-Processing Path")
add_paragraph(doc,
    "Pipeline A operates directly on the raw EEG amplitude values and applies three sequential "
    "transformations: (1) MinMaxScaler normalises each feature to [0, 1] independently, "
    "preventing amplitude differences between EEG channels from dominating the classifier; "
    "(2) a 4th-order Butterworth bandpass filter (0.5–50 Hz, implemented via "
    "scipy.signal.sosfiltfilt) removes power-line noise and DC drift while preserving the "
    "clinically relevant delta (0.5–4 Hz), theta (4–8 Hz), alpha (8–13 Hz), and beta "
    "(13–30 Hz) frequency bands; (3) SelectKBest with the ANOVA F-statistic selects the "
    "top k=50 features, reducing dimensionality while retaining the most discriminative "
    "time-domain points.")

add_subheading(doc, "B.  Pipeline B: Feature-Extraction Path")
add_paragraph(doc,
    "Pipeline B transforms raw signals into a compact statistical feature vector before "
    "applying any scaling, following a fundamentally different data representation strategy. "
    "Step (1) extracts nine hand-crafted statistical features per sample: mean, standard "
    "deviation, skewness, kurtosis, signal energy (sum of squared values), peak absolute "
    "amplitude, minimum value, peak-to-peak range, and zero-crossing rate. This reduces "
    "the 178-dimensional input to a 9-dimensional feature vector. Step (2) applies "
    "StandardScaler (zero mean, unit variance) appropriate for Gaussian-distributed features. "
    "Step (3) applies PCA retaining 8 principal components (~88\\% of variance), further "
    "decorrelating the features and providing additional regularisation.")

add_ieee_table(doc,
    ["Step", "Pipeline A  (Signal Path)", "Pipeline B  (Feature Path)"],
    [
        ["1", "MinMaxScaler  [0, 1]",           "Statistical Feature Extraction  (9 features)"],
        ["2", "Butterworth Bandpass Filter",     "StandardScaler  (zero mean, unit var)"],
        ["3", "SelectKBest (ANOVA F, k=50)",     "PCA  (8 components)"],
        ["Output dim.", "50 features",           "8 features"],
        ["Preserves raw signal?", "Yes",         "No — aggregated statistics only"],
    ],
    "TABLE III.  Preprocessing Pipeline Comparison",
    col_widths=[1.2, 2.9, 2.9],
)

# ─────────────────────────────────────────────────────────────────────────────
# IV. MODEL: LOGISTIC REGRESSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Model: Logistic Regression with Regularisation", "IV.")
add_paragraph(doc,
    "Logistic Regression is selected as the sole classifier for this study to isolate the "
    "effects of preprocessing and regularisation from model complexity. The probability of "
    "a positive (seizure) outcome is modelled as:")

# Equation block
eq_p = doc.add_paragraph()
eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_p.paragraph_format.space_before = Pt(4)
eq_p.paragraph_format.space_after  = Pt(4)
r = eq_p.add_run("P(y=1 | x) = 1 / (1 + exp(−(β₀ + βᵀx)))")
set_font(r, size=11, italic=True)

add_paragraph(doc,
    "Training minimises the regularised cross-entropy loss. Under the unified ElasticNet "
    "formulation used by scikit-learn v1.8, the objective is controlled by a single "
    "l1_ratio parameter ρ ∈ [0, 1]:")

eq_p2 = doc.add_paragraph()
eq_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_p2.paragraph_format.space_before = Pt(4)
eq_p2.paragraph_format.space_after  = Pt(4)
r = eq_p2.add_run(
    "J(β) = (1/m)Σ L(ŷ⁻ⁱ⁾, y⁻ⁱ⁾) + "
    "(λ / 2m)[ρ||β||₁ + (1−ρ)||β||"
    "²₂]"
)
set_font(r, size=11, italic=True)

add_paragraph(doc,
    "Setting ρ=1.0 yields pure L1 (Lasso), which drives coefficients to exactly zero "
    "and implicitly performs feature selection. Setting ρ=0.0 yields pure L2 (Ridge), "
    "which shrinks all coefficients uniformly. ElasticNet (ρ=0.5) interpolates between "
    "the two. All models use the SAGA solver, which supports all three penalty types, with "
    "max_iter=5,000 and C=1.0 (where C=1/λ is the inverse regularisation strength). "
    "The train/test split is 80/20 with stratification to preserve the class ratio.")

# ─────────────────────────────────────────────────────────────────────────────
# V. CLASS IMBALANCE HANDLING
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Class Imbalance Handling", "V.")
add_paragraph(doc,
    "Class imbalance is a critical challenge in seizure detection. We evaluate three "
    "complementary strategies:")

add_subheading(doc, "A.  SMOTE Oversampling")
add_paragraph(doc,
    "Synthetic Minority Oversampling Technique (SMOTE) [6] generates synthetic minority-class "
    "samples by interpolating between existing minority instances in feature space. For UCI, "
    "SMOTE expands the training set from 9,200 to 14,720 samples (balanced 50/50). For "
    "CHB-MIT, the training set grows from 4,320 to 8,608 samples. The number of nearest "
    "neighbours k is clamped to min(5, minority_count − 1) to handle extreme imbalance.")

add_subheading(doc, "B.  Random Undersampling")
add_paragraph(doc,
    "RandomUnderSampler reduces the majority class by randomly discarding samples until "
    "class balance is achieved. For UCI, the training set shrinks from 9,200 to 3,680 "
    "samples. For CHB-MIT (270:1 ratio), the training set reduces to just 32 samples "
    "(16 per class), representing the most aggressive data reduction. This strategy "
    "prioritises recall over accuracy by eliminating the majority-class advantage.")

add_subheading(doc, "C.  Class Weighting")
add_paragraph(doc,
    "Class weighting adjusts the logistic regression loss function by assigning a weight "
    "inversely proportional to class frequency: w_c = n_samples / (n_classes × n_c). "
    "This requires no resampling and preserves the original data distribution, making it "
    "computationally efficient and appropriate when the training set is already small.")

# ─────────────────────────────────────────────────────────────────────────────
# VI. EXPERIMENTAL RESULTS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Experimental Results", "VI.")
add_paragraph(doc,
    "All 36 experiment configurations (2 datasets × 2 pipelines × 3 penalties "
    "× 3 imbalance strategies) were executed. Three primary metrics are reported: "
    "Accuracy, F1-score (the harmonic mean of precision and recall, prioritised due to "
    "class imbalance), and PR-AUC (area under the Precision-Recall curve, particularly "
    "informative under high imbalance). Note that high Accuracy on imbalanced datasets "
    "(e.g., 97.9\\% on CHB-MIT Pipeline A) is misleading and should not be interpreted "
    "as model success.")

add_subheading(doc, "A.  UCI Dataset — Full Results")

# UCI results table
uci_rows = []
for _, row in uci.iterrows():
    uci_rows.append([
        f"{row['pipeline']}",
        row["penalty"].upper(),
        row["imbalance_strategy"],
        f"{row['accuracy']:.4f}",
        f"{row['f1_score']:.4f}",
        f"{row['pr_auc']:.4f}",
    ])
add_ieee_table(doc,
    ["Pipeline", "Penalty", "Imbalance Strategy", "Accuracy", "F1-Score", "PR-AUC"],
    uci_rows,
    "TABLE IV.  Full UCI Experimental Results (18 configurations)",
    col_widths=[1.1, 0.75, 1.35, 0.9, 0.9, 0.9],
)

add_subheading(doc, "B.  CHB-MIT Dataset — Full Results")

chb_rows = []
for _, row in chb.iterrows():
    chb_rows.append([
        f"{row['pipeline']}",
        row["penalty"].upper(),
        row["imbalance_strategy"],
        f"{row['accuracy']:.4f}",
        f"{row['f1_score']:.4f}",
        f"{row['pr_auc']:.4f}",
    ])
add_ieee_table(doc,
    ["Pipeline", "Penalty", "Imbalance Strategy", "Accuracy", "F1-Score", "PR-AUC"],
    chb_rows,
    "TABLE V.  Full CHB-MIT Experimental Results (18 configurations)",
    col_widths=[1.1, 0.75, 1.35, 0.9, 0.9, 0.9],
)

# ─────────────────────────────────────────────────────────────────────────────
# VII. COMPARATIVE ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Comparative Analysis", "VII.")

add_subheading(doc, "Q1.  Does Preprocessing Order Affect Results?")
add_paragraph(doc,
    "Yes — decisively. Preprocessing order is the single largest performance factor in this "
    "study, dwarfing the effect of regularisation choice or imbalance strategy. On the UCI "
    "dataset, Pipeline B achieves a mean F1-score of 0.895 across all 9 penalty/strategy "
    "combinations, compared to 0.371 for Pipeline A—a 141\\% relative improvement on "
    "identical raw data with the same model. This stark contrast arises because Pipeline B's "
    "statistical feature extraction (mean, std, skewness, kurtosis, energy, zero-crossing "
    "rate) captures discriminative temporal structure that raw amplitude values alone "
    "cannot represent efficiently.")

add_figure(doc,
    os.path.join(RESULTS, "analysis_q1_pipeline_comparison.png"),
    "Fig. 1.  Q1: Pipeline A vs. Pipeline B — mean F1-score and PR-AUC per dataset. "
    "Pipeline B dramatically outperforms Pipeline A on UCI; both fail similarly on CHB-MIT "
    "Pipeline A due to the 270:1 class ratio.",
)

add_subheading(doc, "Q2.  Which Regularisation Generalises Best?")

pen_table_rows = []
for penalty in ["l1", "l2", "elasticnet"]:
    sub   = df[df["penalty"] == penalty]
    uci_f1   = sub[sub["dataset"] == "UCI"]["f1_score"].mean()
    uci_auc  = sub[sub["dataset"] == "UCI"]["pr_auc"].mean()
    chb_f1   = sub[sub["dataset"] == "CHB-MIT"]["f1_score"].mean()
    chb_auc  = sub[sub["dataset"] == "CHB-MIT"]["pr_auc"].mean()
    pen_table_rows.append([
        penalty.upper(),
        f"{uci_f1:.4f}", f"{uci_auc:.4f}",
        f"{chb_f1:.4f}", f"{chb_auc:.4f}",
    ])
add_ieee_table(doc,
    ["Penalty", "UCI F1", "UCI PR-AUC", "CHB-MIT F1", "CHB-MIT PR-AUC"],
    pen_table_rows,
    "TABLE VI.  Regularisation Comparison (means across all pipelines and imbalance strategies)",
    col_widths=[0.85, 1.2, 1.2, 1.2, 1.2],
)
add_paragraph(doc,
    "L2 achieves the highest mean F1 on UCI (0.634 vs 0.632 L1, 0.633 ElasticNet) and is "
    "marginally the most stable regulariser across datasets. However, the differences between "
    "L1, L2, and ElasticNet are less than 0.003 in F1 and 0.004 in PR-AUC, confirming that "
    "regularisation type has a secondary effect compared to pipeline choice.")

add_figure(doc,
    os.path.join(RESULTS, "analysis_q2q3_regularization_comparison.png"),
    "Fig. 2.  Q2/Q3: L1 vs. L2 vs. ElasticNet — F1-score and PR-AUC. "
    "Differences are minimal, suggesting the preprocessing pipeline dominates.",
)

add_subheading(doc, "Q3.  Does ElasticNet Consistently Outperform L1 and L2?")
add_paragraph(doc,
    "No. ElasticNet does not consistently outperform either pure penalty. On UCI, ElasticNet "
    "ranks second in F1 behind L2 but ahead of L1. On CHB-MIT, the ordering is dataset- and "
    "pipeline-dependent. This result is consistent with the theoretical expectation: ElasticNet's "
    "advantage emerges primarily in correlated high-dimensional feature spaces [7]. In our "
    "setting, after PCA (Pipeline B) or SelectKBest (Pipeline A), feature correlations are "
    "substantially reduced, diminishing ElasticNet's advantage.")

add_subheading(doc, "Q4.  How Does Imbalance Handling Interact with Regularisation?")
add_paragraph(doc,
    "The interaction is dataset-dependent. On UCI (4:1 ratio), all three strategies perform "
    "similarly because the imbalance is moderate and the dataset is large enough for all "
    "approaches to succeed. On CHB-MIT (270:1 ratio), the interaction is significant: random "
    "undersampling achieves the highest PR-AUC (0.625 with L1, Pipeline B) despite reducing "
    "the training set to 32 samples. This counter-intuitive result reflects the precision-recall "
    "tradeoff: undersampling forces the model to prioritise minority class recall, which is "
    "captured by PR-AUC. SMOTE achieves slightly higher F1 (0.216 vs 0.190 for UnderSample "
    "with L1) but lower PR-AUC, indicating more conservative recall behaviour.")

add_figure(doc,
    os.path.join(RESULTS, "analysis_q4_imbalance_comparison.png"),
    "Fig. 3.  Q4: SMOTE vs. UnderSample vs. ClassWeight — F1-score and PR-AUC per dataset. "
    "On the extreme CHB-MIT imbalance, UnderSample yields the highest PR-AUC.",
)

add_figure(doc,
    os.path.join(RESULTS, "analysis_precision_recall_tradeoff.png"),
    "Fig. 4.  Precision-Recall tradeoff scatter by imbalance strategy. "
    "Upper-right = optimal (high F1 and high PR-AUC simultaneously).",
)

add_subheading(doc, "Summary Heatmap")
add_paragraph(doc,
    "Fig. 5 presents a heatmap of F1-scores across all 36 configurations. The top-left "
    "cluster (CHB-MIT, Pipeline A) is uniformly near zero, demonstrating underfitting on "
    "an impossibly imbalanced signal-space representation. The bottom-right cluster "
    "(UCI, Pipeline B) is uniformly high (0.89–0.90), showing that Pipeline B with any "
    "regularisation and any imbalance strategy is robust on this dataset.")

add_figure(doc,
    os.path.join(RESULTS, "analysis_summary_heatmap.png"),
    "Fig. 5.  F1-Score summary heatmap across all 36 configurations "
    "(dataset × pipeline vs. penalty × imbalance strategy).",
    width=6.2,
)

# ─────────────────────────────────────────────────────────────────────────────
# VIII. OVERFITTING AND UNDERFITTING
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Overfitting and Underfitting Analysis", "VIII.")
add_paragraph(doc,
    "To explicitly demonstrate both failure modes, we conduct two experiments on the UCI "
    "dataset using Pipeline A features.")

add_subheading(doc, "A.  Explicit Scenario Comparison")

ov_rows = [
    ["Underfitting  (C=0.001, k=5 features)",  "0.001",  "5",   "0.369", "0.369", "0.000"],
    ["Normal        (C=1.0,  k=50 features)",  "1.0",   "50",   "0.375", "0.370", "0.005"],
    ["Overfitting   (C=1000, k=178 features)", "1000",  "178",  "0.990", "0.366", "0.624"],
]
add_ieee_table(doc,
    ["Scenario", "C", "Features k", "Train F1", "Val F1", "Gap"],
    ov_rows,
    "TABLE VII.  Overfitting and Underfitting Scenarios (UCI, Pipeline A, L2)",
    col_widths=[2.6, 0.55, 0.9, 0.8, 0.7, 0.6],
)

add_paragraph(doc,
    "At C=0.001 (strong regularisation) with only 5 features, both training and validation "
    "F1 are nearly identical and low (~0.369), the hallmark of underfitting: the model "
    "lacks the capacity and feature information to separate the classes. At C=1000 (negligible "
    "regularisation) with all 178 features, training F1 reaches 0.990 while validation F1 "
    "drops to 0.366—a gap of 0.624 points—clearly demonstrating overfitting: the model "
    "memorises the training data but generalises poorly.")

add_figure(doc,
    os.path.join(RESULTS, "analysis_overfit_underfit_scenarios.png"),
    "Fig. 6.  Train vs. validation F1 across three explicit scenarios showing underfitting "
    "(small gap, both low), normal fitting, and overfitting (large train-val gap).",
)

add_subheading(doc, "B.  Regularisation Path (C Sweep)")
add_paragraph(doc,
    "Fig. 7 shows the continuous regularisation path obtained by sweeping C from 10⁻⁴ "
    "to 10³ on Pipeline A features (k=50). The shaded blue region (low C) marks the "
    "underfitting zone where both train and validation F1 are suppressed by excessive "
    "regularisation. The shaded red region (high C) marks the overfitting zone where the "
    "train-validation gap grows. The optimal C lies between these extremes, confirming that "
    "regularisation strength must be tuned even for linear classifiers.")

add_figure(doc,
    os.path.join(RESULTS, "reg_path_UCI_Pipeline_A_50feat.png"),
    "Fig. 7.  Regularisation path: train vs. validation F1 as C varies (log scale). "
    "Blue zone = underfitting; red zone = overfitting.",
)

add_subheading(doc, "C.  Learning Curves")
add_paragraph(doc,
    "Fig. 8 shows the learning curve for Pipeline B + L2 Logistic Regression on UCI, "
    "computed via 5-fold stratified cross-validation over 10 training set sizes from 50\\% "
    "to 100\\% of the training data. The training and cross-validation F1 curves converge "
    "to approximately 0.89 as the training set grows, with a small and diminishing gap. "
    "This indicates that: (1) the model is not significantly overfitting at full data size; "
    "(2) additional data beyond the available 9,200 training samples would provide only "
    "marginal benefit; (3) the chosen pipeline and regularisation generalise well.")

add_figure(doc,
    os.path.join(RESULTS, "learning_curve_UCI_Pipeline_B_L2.png"),
    "Fig. 8.  Learning curves for Pipeline B + L2 on UCI. Shaded regions show ±1 std. "
    "across 5 CV folds. Converging curves indicate good bias-variance balance.",
)

# ─────────────────────────────────────────────────────────────────────────────
# IX. REGULARISATION STUDY: SPARSITY
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Regularisation Study: Sparsity Analysis", "IX.")
add_paragraph(doc,
    "To analyse the feature-selection effect of regularisation, we fit L1, ElasticNet, and L2 "
    "logistic regression models directly on the 178-dimensional normalised and denoised UCI "
    "feature space (MinMaxScaler + Butterworth filter, no SelectKBest). Table VIII summarises "
    "the coefficient sparsity.")

add_ieee_table(doc,
    ["Regularisation", "l1_ratio (ρ)", "Non-Zero Coefficients", "Zero Coefficients", "Sparsity (%)"],
    [
        ["L1  (Lasso)",  "1.0",  "46 / 178",  "132 / 178", "74.2%"],
        ["ElasticNet",   "0.5", "109 / 178",   "69 / 178", "38.8%"],
        ["L2  (Ridge)",  "0.0", "178 / 178",    "0 / 178",  "0.0%"],
    ],
    "TABLE VIII.  Coefficient Sparsity at C=1.0 (UCI, all 178 features)",
    col_widths=[1.4, 1.1, 1.6, 1.6, 1.1],
)

add_paragraph(doc,
    "L1 regularisation drives 132 of 178 coefficients to exactly zero, selecting only "
    "46 features (25.8\\% of the feature space) as discriminative. ElasticNet retains "
    "109 features—a middle ground that inherits some of L1's sparse structure while "
    "remaining more stable than pure L1. L2 retains all 178 features with small, "
    "non-zero weights. This sparsity difference has practical implications: L1 and "
    "ElasticNet are self-interpreting in that the non-zero coefficients directly indicate "
    "which time-domain EEG points are most discriminative for seizure classification.")

add_figure(doc,
    os.path.join(RESULTS, "sparsity_UCI_all_features.png"),
    "Fig. 9.  Sorted absolute coefficient magnitudes (log scale) for L1, ElasticNet, "
    "and L2. L1's step-function profile shows 132 zero coefficients; "
    "L2's gradual decay shows all 178 features receive non-zero weight.",
)

# ─────────────────────────────────────────────────────────────────────────────
# X. DISCUSSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Discussion", "X.")
add_paragraph(doc,
    "The results of this study yield four principal insights relevant to the broader "
    "seizure detection literature:")

add_paragraph(doc,
    "1) Preprocessing dominates performance. The 141\\% F1 improvement from Pipeline A to "
    "Pipeline B on UCI—with an identical model—demonstrates that feature engineering "
    "is more impactful than regularisation choice for linear classifiers. Practitioners "
    "should prioritise feature extraction over hyperparameter tuning.")

add_paragraph(doc,
    "2) Accuracy is an unreliable metric under imbalance. CHB-MIT Pipeline A achieves "
    "97.9\\% accuracy while predicting zero seizures (F1=0.000). PR-AUC should be the "
    "primary metric for imbalanced medical datasets.")

add_paragraph(doc,
    "3) Extreme imbalance requires aggressive intervention. The 270:1 CHB-MIT ratio "
    "renders standard SMOTE insufficient when applied to high-dimensional raw features. "
    "Only after dimensionality reduction (Pipeline B) does SMOTE achieve meaningful "
    "recall, with undersampling providing superior PR-AUC.")

add_paragraph(doc,
    "4) ElasticNet's advantage is context-dependent. The theoretical superiority of "
    "ElasticNet over L1/L2 for correlated features [7] does not materialise after PCA "
    "decorrelates the feature space. Regular L2 regularisation is a strong, stable "
    "baseline for post-PCA classification tasks.")

# ─────────────────────────────────────────────────────────────────────────────
# XI. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Conclusion", "XI.")
add_paragraph(doc,
    "This paper presented a comprehensive end-to-end study of epileptic seizure detection "
    "using Logistic Regression across two EEG datasets, two preprocessing pipelines, three "
    "regularisation strategies, and three class-imbalance handling techniques. The key "
    "finding is that the choice of preprocessing pipeline is the dominant performance factor "
    "for linear classifiers: statistical feature extraction (Pipeline B) outperforms "
    "direct signal processing (Pipeline A) by 141\\% in F1-score on the UCI dataset. "
    "Regularisation type and imbalance handling strategy have secondary but measurable "
    "effects, particularly under extreme imbalance. L1 regularisation provides built-in "
    "feature selection (74.2\\% sparsity), and random undersampling outperforms SMOTE in "
    "PR-AUC under a 270:1 class imbalance ratio. Future work should extend this study to "
    "non-linear models (Support Vector Machines, Random Forests, deep convolutional "
    "networks) and include the Kaggle iEEG dataset for intracranial EEG validation.")

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
add_ruled_line(doc)
add_heading(doc, "References", "", size=10, space_before=6)

refs = [
    "[1] World Health Organization, \"Epilepsy,\" WHO Fact Sheet, Feb. 2024. [Online]. "
    "Available: https://www.who.int/news-room/fact-sheets/detail/epilepsy",

    "[2] R. S. Fisher et al., \"ILAE Official Report: A practical clinical definition of "
    "epilepsy,\" Epilepsia, vol. 55, no. 4, pp. 475–482, Apr. 2014.",

    "[3] R. G. Andrzejak et al., \"Indications of nonlinear deterministic and finite "
    "dimensional structures in time series of brain electrical activity,\" Physical Review E, "
    "vol. 64, no. 6, 2001. [Dataset] UCI Machine Learning Repository, ID 388.",

    "[4] A. L. Goldberger et al., \"PhysioBank, PhysioToolkit, and PhysioNet: Components of "
    "a New Research Resource for Complex Physiologic Signals,\" Circulation, vol. 101, "
    "no. 23, pp. e215–e220, 2000. CHB-MIT Scalp EEG Database v1.0.0.",

    "[5] Kaggle / UPenn and Mayo Clinic, \"Seizure Detection,\" Kaggle Competition, 2014. "
    "[Online]. Available: https://www.kaggle.com/c/seizure-detection",

    "[6] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, \"SMOTE: Synthetic "
    "Minority Over-sampling Technique,\" J. Artif. Intell. Res., vol. 16, pp. 321–357, 2002.",

    "[7] H. Zou and T. Hastie, \"Regularization and Variable Selection via the Elastic Net,\" "
    "J. Royal Stat. Soc. Series B, vol. 67, no. 2, pp. 301–320, 2005.",

    "[8] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" J. Machine "
    "Learning Research, vol. 12, pp. 2825–2830, 2011.",

    "[9] G. Lemaître, F. Nogueira, and C. K. Aridas, \"Imbalanced-learn: A Python Toolbox "
    "to Tackle the Curse of Imbalanced Datasets in Machine Learning,\" J. Machine Learning "
    "Research, vol. 18, no. 17, pp. 1–5, 2017.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(ref)
    set_font(r, size=9)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
print(f"Size:  {os.path.getsize(OUT_FILE) / 1024:.1f} KB")
